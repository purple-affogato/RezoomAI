from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt
from supabase import Client

# to be completed
def get_docx(client:Client, text:str, name, contacts):
    doc = Document()
    name_line = doc.add_paragraph(name)
    name_line.style.font.size = Pt(12)
    name_line.style.font.name = "Arial"
    name_line.style.font.bold = True
    name_line.paragraph_format.space_after = 0
    contacts_line = doc.add_paragraph(contacts)
    contacts_line.style.font.size = Pt(12)
    contacts_line.style.font.name = "Arial"
    contacts_line.paragraph_format.space_after = 0
    para = doc.add_paragraph(text)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    para.style.font.size = Pt(12)
    para.style.font.name = "Arial"
    para.paragraph_format.space_after = 0
    file_path = "resume.docx"
    doc.save(file_path)
    # upload to bucket
    with open('./'+file_path, 'rb') as f:
        response = client.storage.from_("bucky").upload(
        file=f,
        path=file_path,
        file_options={"cache-control": "3600", "upsert": "false"},
    )
        
    # send docx
    response = client.storage.from_("bucky").get_public_url(file_path)
    return response
    